# -*- coding: UTF-8 -*-
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm
from pypinyin import pinyin


def add_pinyin(cell, i):
    cell.text = py[i]
    cell.width = Cm(1.56)
    paragraph = cell.paragraphs[0]
    paragraph.style.font.name = 'Arial'
    paragraph.style.font.size = Pt(14)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_picture('pic.png')


def add_new_table(document, py_len):
    table = document.add_table(rows=1, cols=12)
    for i in range(0, py_len):
        cell = table.rows[0].cells[i]
        add_pinyin(cell, i)
    table.rows[0].cells[py_len].width = Cm(1)


f = open("words.txt", encoding='UTF-8')
line = f.readline()
words = []
while line:
    words.append(line.replace('\n', ''))
    line = f.readline()
f.close()

num = 0
document = Document()
sections = document.sections
section = sections[0]
section.left_margin = Cm(1.9)
section.right_margin = Cm(1.9)
for w in words:
    py = pinyin(w)
    py_len = py.__len__()
    if num == 0:
        add_new_table(document, py_len)
        num = num + py_len
    elif num + py_len < 12:
        tables = document.tables
        table = tables[len(tables) - 1]
        for i in range(0, py_len):
            cell = table.rows[0].cells[num + i + 1]
            add_pinyin(cell, i)
        try:
            table.rows[0].cells[num + py_len + 1].width = Cm(1)
        except IndexError:
            pass
        num = num + py_len + 1
    else:
        num = 0
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run("")
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(0)
        add_new_table(document, py_len)
        num = num + py_len

document.save("D:\\看拼音写汉字.docx")
